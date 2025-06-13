import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
from datetime import datetime, timedelta
import sys
import os
import json
import numpy as np
from plotly.subplots import make_subplots
from sklearn.metrics import mean_squared_error, mean_absolute_error, r2_score
import torch
from collections import defaultdict
from scipy import stats
from scipy.signal import find_peaks
from scipy.stats import norm
from scipy.optimize import minimize
import cvxopt
from cvxopt import matrix, solvers
from enum import Enum
import uuid
import plotly.figure_factory as ff
from plotly.subplots import make_subplots
import seaborn as sns
import matplotlib.pyplot as plt
from matplotlib.colors import LinearSegmentedColormap
import networkx as nx
from wordcloud import WordCloud
import folium
from folium.plugins import HeatMap
import altair as alt
import vega_datasets
import pydeck as pdk
import streamlit.components.v1 as components
import base64
from io import BytesIO
import warnings
import csv
import xlsxwriter
from docx import Document
from docx.shared import Inches
import pdfkit
import jinja2
import os
warnings.filterwarnings('ignore')

# Add both project root and trading directory to Python path
project_root = os.path.dirname(os.path.abspath(__file__))
sys.path.insert(0, project_root)
sys.path.insert(0, os.path.join(project_root, 'trading'))

# Import our custom modules
from trading.strategies import StrategyManager
from trading.backtesting import BacktestEngine
from trading.risk import RiskManager
from trading.portfolio import PortfolioManager
from trading.analysis import MarketAnalyzer
from trading.optimization import Optimizer
from trading.llm import LLMInterface
from trading.feature_engineering import FeatureEngineer
from trading.evaluation import ModelEvaluator
from trading.knowledge_base.trading_rules import TradingRules
from trading.market.market_data import MarketData
from trading.execution.execution_engine import ExecutionEngine
from trading.market.market_indicators import MarketIndicators
from trading.models.lstm_model import LSTMForecaster
from trading.nlp.llm_processor import LLMProcessor

# Set page config
st.set_page_config(
    page_title="Evolve Clean Trading Dashboard",
    page_icon="📈",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Initialize session state variables first
if 'use_api' not in st.session_state:
    st.session_state.use_api = False
if 'api_key' not in st.session_state:
    st.session_state.api_key = ""

# Initialize components
if 'portfolio' not in st.session_state:
    st.session_state.portfolio = PortfolioManager()
if 'risk_manager' not in st.session_state:
    st.session_state.risk_manager = RiskManager(pd.Series(dtype=float))
if 'strategy_manager' not in st.session_state:
    st.session_state.strategy_manager = StrategyManager()
if 'backtest_engine' not in st.session_state:
    st.session_state.backtest_engine = BacktestEngine(pd.DataFrame())
if 'market_analyzer' not in st.session_state:
    st.session_state.market_analyzer = MarketAnalyzer()
if 'optimizer' not in st.session_state:
    st.session_state.optimizer = Optimizer(state_dim=10, action_dim=5)
if 'feature_engineer' not in st.session_state:
    st.session_state.feature_engineer = FeatureEngineer()
if 'model_evaluator' not in st.session_state:
    st.session_state.model_evaluator = ModelEvaluator()

# Initialize LLM last, after all other components
if 'llm' not in st.session_state:
    api_key = st.session_state.api_key if st.session_state.use_api else None
    st.session_state.llm = LLMInterface(api_key=api_key)

# Initialize TradingRules
if 'trading_rules' not in st.session_state:
    st.session_state.trading_rules = TradingRules()
    st.session_state.trading_rules.load_rules_from_json('trading/knowledge_base/trading_rules.json')

# Initialize MarketData
if 'market_data' not in st.session_state:
    st.session_state.market_data = MarketData()

# Initialize ExecutionEngine
if 'execution_engine' not in st.session_state:
    st.session_state.execution_engine = ExecutionEngine()

# Initialize MarketIndicators
if 'market_indicators' not in st.session_state:
    st.session_state.market_indicators = MarketIndicators()

# Initialize LSTM model
if 'lstm_model' not in st.session_state:
    lstm_config = {
        'input_size': 5,  # Number of features
        'hidden_size': 64,
        'num_layers': 2,
        'dropout': 0.2,
        'sequence_length': 10,
        'feature_columns': ['open', 'high', 'low', 'close', 'volume'],
        'target_column': 'close',
        'use_attention': True,
        'use_batch_norm': True,
        'max_sequence_length': 100,
        'max_batch_size': 32
    }
    st.session_state.lstm_model = LSTMForecaster(lstm_config)
    st.session_state.model_metrics = {}

# Initialize MarketAnalyzer
if 'market_analyzer' not in st.session_state:
    analyzer_config = {
        'trend_threshold': 0.02,  # 2% threshold for trend strength
        'volatility_window': 252,  # 1 year for volatility calculation
        'correlation_threshold': 0.7  # Strong correlation threshold
    }
    st.session_state.market_analyzer = MarketAnalyzer(analyzer_config)

# Initialize LLM Processor with enhanced configuration
if 'llm_processor' not in st.session_state:
    st.session_state.llm_processor = LLMProcessor()
    st.session_state.news_cache = {}
    st.session_state.event_cache = {}
    st.session_state.entity_cache = {}
    st.session_state.sentiment_signals = {}
    st.session_state.signal_history = {}
    st.session_state.signal_performance = {}
    st.session_state.market_filters = {}
    st.session_state.filter_thresholds = {}
    st.session_state.market_regime = {}
    st.session_state.regime_transitions = {}
    st.session_state.regime_statistics = {}
    st.session_state.signal_patterns = {}
    st.session_state.backtest_results = {}
    st.session_state.portfolio_weights = {}
    st.session_state.risk_metrics = {}
    st.session_state.optimization_results = {}
    st.session_state.orders = {}
    st.session_state.execution_strategies = {}
    st.session_state.order_history = {}
    st.session_state.positions = {}
    st.session_state.risk_limits = {}
    st.session_state.position_limits = {}
    st.session_state.portfolio_risk = {
        'total_exposure': 0.0,
        'net_exposure': 0.0,
        'gross_exposure': 0.0,
        'portfolio_var': 0.0,
        'portfolio_volatility': 0.0,
        'portfolio_beta': 0.0,
        'portfolio_correlation': 0.0,
        'portfolio_sharpe': 0.0,
        'portfolio_sortino': 0.0,
        'portfolio_drawdown': 0.0,
        'portfolio_leverage': 0.0,
        'sector_exposure': defaultdict(float),
        'factor_exposure': defaultdict(float),
        'risk_contributions': defaultdict(float),
        'last_update': datetime.now()
    }
    st.session_state.portfolio_limits = {
        'max_total_exposure': 1.0,
        'max_net_exposure': 0.3,
        'max_gross_exposure': 1.5,
        'max_portfolio_var': 0.02,
        'max_portfolio_volatility': 0.15,
        'max_portfolio_beta': 1.2,
        'max_portfolio_correlation': 0.7,
        'min_portfolio_sharpe': 0.8,
        'min_portfolio_sortino': 1.0,
        'max_portfolio_drawdown': 0.1,
        'max_portfolio_leverage': 1.0,
        'max_sector_exposure': 0.3,
        'max_factor_exposure': 0.4,
        'max_risk_contribution': 0.2,
        'last_update': datetime.now()
    }
    st.session_state.portfolio_optimization = {
        'target_weights': {},
        'current_weights': {},
        'rebalance_threshold': 0.05,
        'optimization_method': 'risk_parity',
        'risk_free_rate': 0.02,
        'optimization_constraints': {},
        'last_optimization': None,
        'optimization_history': []
    }
    st.session_state.execution_management = {
        'execution_strategies': {
            'TWAP': {
                'time_horizon': 3600,  # 1 hour
                'num_slices': 12,
                'min_slice_size': 100,
                'max_slice_size': 1000,
                'price_improvement': 0.0001
            },
            'VWAP': {
                'volume_profile': 'U',
                'min_participation': 0.1,
                'max_participation': 0.3,
                'price_improvement': 0.0002
            },
            'POV': {
                'participation_rate': 0.2,
                'min_participation': 0.1,
                'max_participation': 0.3,
                'price_improvement': 0.0003
            },
            'IS': {
                'urgency': 'NORMAL',
                'min_participation': 0.05,
                'max_participation': 0.2,
                'price_improvement': 0.0004
            }
        },
        'order_management': {
            'max_orders_per_symbol': 5,
            'min_order_interval': 300,  # 5 minutes
            'max_slippage': 0.001,
            'min_liquidity': 1000000,
            'max_spread': 0.002,
            'price_improvement_target': 0.0005,
            'execution_cost_target': 0.001,
            'market_impact_target': 0.002
        },
        'execution_metrics': {
            'total_orders': 0,
            'filled_orders': 0,
            'cancelled_orders': 0,
            'average_slippage': 0.0,
            'average_execution_time': 0.0,
            'average_price_improvement': 0.0,
            'average_market_impact': 0.0,
            'execution_cost': 0.0,
            'last_update': datetime.now()
        }
    }
    st.session_state.market_impact = {
        'impact_models': {
            'square_root': {
                'alpha': 0.1,
                'beta': 0.5,
                'gamma': 0.1
            },
            'linear': {
                'alpha': 0.2,
                'beta': 0.3,
                'gamma': 0.1
            },
            'power_law': {
                'alpha': 0.15,
                'beta': 0.4,
                'gamma': 0.2
            }
        },
        'market_conditions': {
            'volatility_factor': 1.0,
            'liquidity_factor': 1.0,
            'spread_factor': 1.0,
            'volume_factor': 1.0,
            'momentum_factor': 1.0,
            'last_update': datetime.now()
        },
        'impact_metrics': {
            'temporary_impact': 0.0,
            'permanent_impact': 0.0,
            'total_impact': 0.0,
            'impact_cost': 0.0,
            'opportunity_cost': 0.0,
            'last_update': datetime.now()
        }
    }
    st.session_state.adaptive_execution = {
        'execution_parameters': {
            'urgency_level': 'NORMAL',
            'participation_rate': 0.2,
            'time_horizon': 3600,
            'price_improvement_target': 0.0005,
            'market_impact_target': 0.002,
            'execution_cost_target': 0.001
        },
        'adaptation_rules': {
            'volatility_threshold': 0.2,
            'liquidity_threshold': 1000000,
            'spread_threshold': 0.002,
            'volume_threshold': 100000,
            'momentum_threshold': 0.1
        },
        'adaptation_history': [],
        'last_adaptation': None
    }
    st.session_state.performance_analytics = {
        'returns': {
            'daily': defaultdict(list),
            'weekly': defaultdict(list),
            'monthly': defaultdict(list),
            'yearly': defaultdict(list)
        },
        'risk_metrics': {
            'sharpe_ratio': defaultdict(list),
            'sortino_ratio': defaultdict(list),
            'max_drawdown': defaultdict(list),
            'var_95': defaultdict(list),
            'expected_shortfall': defaultdict(list),
            'beta': defaultdict(list),
            'correlation': defaultdict(list)
        },
        'execution_metrics': {
            'slippage': defaultdict(list),
            'execution_time': defaultdict(list),
            'price_improvement': defaultdict(list),
            'market_impact': defaultdict(list),
            'execution_cost': defaultdict(list)
        },
        'signal_metrics': {
            'accuracy': defaultdict(list),
            'win_rate': defaultdict(list),
            'profit_factor': defaultdict(list),
            'average_return': defaultdict(list),
            'sharpe_ratio': defaultdict(list)
        },
        'portfolio_metrics': {
            'total_return': [],
            'annualized_return': [],
            'volatility': [],
            'sharpe_ratio': [],
            'sortino_ratio': [],
            'max_drawdown': [],
            'var_95': [],
            'expected_shortfall': [],
            'beta': [],
            'correlation': [],
            'information_ratio': [],
            'tracking_error': [],
            'active_share': [],
            'concentration': [],
            'turnover': []
        },
        'regime_metrics': {
            'regime_duration': defaultdict(list),
            'regime_returns': defaultdict(list),
            'regime_volatility': defaultdict(list),
            'regime_sharpe': defaultdict(list),
            'regime_transitions': defaultdict(list)
        },
        'last_update': datetime.now()
    }
    st.session_state.reporting = {
        'performance_summary': {
            'total_return': 0.0,
            'annualized_return': 0.0,
            'volatility': 0.0,
            'sharpe_ratio': 0.0,
            'sortino_ratio': 0.0,
            'max_drawdown': 0.0,
            'win_rate': 0.0,
            'profit_factor': 0.0,
            'average_trade': 0.0,
            'largest_win': 0.0,
            'largest_loss': 0.0,
            'average_win': 0.0,
            'average_loss': 0.0,
            'recovery_factor': 0.0,
            'calmar_ratio': 0.0,
            'omega_ratio': 0.0,
            'tail_ratio': 0.0,
            'value_at_risk': 0.0,
            'expected_shortfall': 0.0,
            'last_update': datetime.now()
        },
        'risk_metrics': {
            'var_95': 0.0,
            'var_99': 0.0,
            'expected_shortfall_95': 0.0,
            'expected_shortfall_99': 0.0,
            'beta': 0.0,
            'correlation': 0.0,
            'tracking_error': 0.0,
            'information_ratio': 0.0,
            'active_share': 0.0,
            'concentration': 0.0,
            'turnover': 0.0,
            'last_update': datetime.now()
        },
        'execution_metrics': {
            'total_orders': 0,
            'filled_orders': 0,
            'cancelled_orders': 0,
            'average_slippage': 0.0,
            'average_execution_time': 0.0,
            'average_price_improvement': 0.0,
            'average_market_impact': 0.0,
            'execution_cost': 0.0,
            'last_update': datetime.now()
        },
        'signal_metrics': {
            'total_signals': 0,
            'accurate_signals': 0,
            'win_rate': 0.0,
            'profit_factor': 0.0,
            'average_return': 0.0,
            'sharpe_ratio': 0.0,
            'last_update': datetime.now()
        },
        'regime_metrics': {
            'current_regime': None,
            'regime_duration': 0.0,
            'regime_return': 0.0,
            'regime_volatility': 0.0,
            'regime_sharpe': 0.0,
            'last_update': datetime.now()
        },
        'visualization_settings': {
            'chart_style': 'plotly',
            'color_scheme': 'default',
            'chart_height': 400,
            'chart_width': 800,
            'show_grid': True,
            'show_legend': True,
            'show_annotations': True,
            'last_update': datetime.now()
        }
    }

# Sidebar Navigation
st.sidebar.title("Evolve Clean")

# Add prompt input to sidebar
st.sidebar.subheader("AI Assistant")
user_prompt = st.sidebar.text_area(
    "Ask me anything about trading",
    height=150,
    help="I can help with market analysis, trading strategies, risk assessment, portfolio management, and more."
)

if user_prompt:
    # Process the prompt using LLM
    response = st.session_state.llm.process_prompt(user_prompt)
    
    # Display the response in the sidebar
    st.sidebar.markdown("---")
    st.sidebar.subheader("Response")
    st.sidebar.write(response.get('content', response.get('error', 'Processing...')))
    
    # Automatically configure system based on response
    if 'actions' in response:
        for action in response['actions']:
            if action['type'] == 'update_chart':
                # Update relevant charts
                if 'chart_type' in action:
                    if action['chart_type'] == 'market_analysis':
                        st.session_state.market_analyzer.update_charts(action['data'])
                    elif action['chart_type'] == 'portfolio':
                        st.session_state.portfolio.update_charts(action['data'])
                    elif action['chart_type'] == 'risk':
                        st.session_state.risk_manager.update_charts(action['data'])
            
            elif action['type'] == 'update_metrics':
                # Update relevant metrics
                if 'metrics' in action:
                    if 'portfolio' in action['metrics']:
                        st.session_state.portfolio.update_metrics(action['metrics']['portfolio'])
                    if 'risk' in action['metrics']:
                        st.session_state.risk_manager.update_metrics(action['metrics']['risk'])
                    if 'strategy' in action['metrics']:
                        st.session_state.strategy_manager.update_metrics(action['metrics']['strategy'])
            
            elif action['type'] == 'show_analysis':
                # Show detailed analysis
                if 'analysis_type' in action:
                    if action['analysis_type'] == 'market':
                        st.session_state.market_analyzer.show_analysis(action['data'])
                    elif action['analysis_type'] == 'risk':
                        st.session_state.risk_manager.show_analysis(action['data'])
                    elif action['analysis_type'] == 'portfolio':
                        st.session_state.portfolio.show_analysis(action['data'])
            
            elif action['type'] == 'configure_strategy':
                # Configure trading strategy
                if 'strategy_params' in action:
                    st.session_state.strategy_manager.configure(action['strategy_params'])
            
            elif action['type'] == 'update_portfolio':
                # Update portfolio allocation
                if 'allocation' in action:
                    st.session_state.portfolio.update_allocation(action['allocation'])
            
            elif action['type'] == 'run_backtest':
                # Run backtest with specified parameters
                if 'backtest_params' in action:
                    st.session_state.backtest_engine.run_backtest(action['backtest_params'])
            
            elif action['type'] == 'optimize_portfolio':
                # Optimize portfolio
                if 'optimization_params' in action:
                    st.session_state.optimizer.optimize(action['optimization_params'])
            
            elif action['type'] == 'update_models':
                # Update ML models
                if 'model_params' in action:
                    st.session_state.model_evaluator.update_models(action['model_params'])
            
            elif action['type'] == 'generate_features':
                # Generate new features
                if 'feature_params' in action:
                    st.session_state.feature_engineer.generate_features(action['feature_params'])
            
            elif action['type'] == 'apply_rule':
                # Apply a trading rule
                if 'rule_params' in action:
                    st.session_state.trading_rules.apply_rule(action['rule_params'])
            
            elif action['type'] == 'fetch_market_data':
                # Fetch market data
                if 'symbol' in action:
                    st.session_state.market_data.fetch_data(
                        action['symbol'],
                        action.get('start_date', (datetime.now() - timedelta(days=365)).strftime('%Y-%m-%d')),
                        action.get('end_date', datetime.now().strftime('%Y-%m-%d'))
                    )
            
            elif action['type'] == 'analyze_market':
                symbol = action.get('symbol', 'AAPL')
                market_symbol = action.get('market_symbol', '^GSPC')  # S&P 500 as default market
                
                # Get historical data
                data = st.session_state.market_data.get_historical_data(symbol, window=100)
                market_data_df = st.session_state.market_data.get_historical_data(market_symbol, window=100)
                
                if data is not None and not data.empty and market_data_df is not None and not market_data_df.empty:
                    # Analyze market conditions
                    market_conditions = st.session_state.market_analyzer.analyze_market_conditions(data, market_data_df)
                    
                    # Calculate advanced market metrics
                    volatility = market_conditions['volatility']['volatility']
                    trend_strength = abs(market_conditions['trend']['trend_strength'])
                    correlation = abs(market_conditions['correlation']['correlation'])
                    
                    # Calculate additional technical indicators
                    rsi = st.session_state.market_indicators.calculate_rsi(data['Close'], window=14)
                    macd = st.session_state.market_indicators.calculate_macd(data['Close'])
                    bollinger_bands = st.session_state.market_indicators.calculate_bollinger_bands(data['Close'])
                    
                    # Calculate market regime with enhanced detection
                    regime_scores = {
                        'HIGH_VOLATILITY': 0,
                        'TRENDING': 0,
                        'CORRELATED': 0,
                        'NORMAL': 0
                    }
                    
                    # Volatility regime scoring
                    if volatility > 0.03:
                        regime_scores['HIGH_VOLATILITY'] += 2
                    elif volatility > 0.02:
                        regime_scores['HIGH_VOLATILITY'] += 1
                    
                    # Trend regime scoring
                    if trend_strength > 0.7:
                        regime_scores['TRENDING'] += 2
                    elif trend_strength > 0.5:
                        regime_scores['TRENDING'] += 1
                    
                    # Correlation regime scoring
                    if correlation > 0.8:
                        regime_scores['CORRELATED'] += 2
                    elif correlation > 0.6:
                        regime_scores['CORRELATED'] += 1
                    
                    # Technical indicator scoring
                    if rsi.iloc[-1] > 70 or rsi.iloc[-1] < 30:
                        regime_scores['HIGH_VOLATILITY'] += 1
                    
                    if abs(macd['macd'].iloc[-1]) > abs(macd['signal'].iloc[-1]):
                        regime_scores['TRENDING'] += 1
                    
                    if data['Close'].iloc[-1] > bollinger_bands['upper'].iloc[-1] or \
                       data['Close'].iloc[-1] < bollinger_bands['lower'].iloc[-1]:
                        regime_scores['HIGH_VOLATILITY'] += 1
                    
                    # Determine dominant regime
                    regime = max(regime_scores.items(), key=lambda x: x[1])[0]
                    regime_confidence = regime_scores[regime] / sum(regime_scores.values())
                    
                    # Store market regime with enhanced information
                    st.session_state.market_regime[symbol] = {
                        'regime': regime,
                        'confidence': regime_confidence,
                        'scores': regime_scores,
                        'timestamp': datetime.now()
                    }
                    
                    # Track regime transitions
                    if symbol not in st.session_state.regime_transitions:
                        st.session_state.regime_transitions[symbol] = []
                    
                    if st.session_state.regime_transitions[symbol]:
                        last_regime = st.session_state.regime_transitions[symbol][-1]['regime']
                        if last_regime != regime:
                            st.session_state.regime_transitions[symbol].append({
                                'from_regime': last_regime,
                                'to_regime': regime,
                                'timestamp': datetime.now(),
                                'confidence': regime_confidence
                            })
                    else:
                        st.session_state.regime_transitions[symbol].append({
                            'from_regime': None,
                            'to_regime': regime,
                            'timestamp': datetime.now(),
                            'confidence': regime_confidence
                        })
                    
                    # Calculate regime statistics
                    if symbol not in st.session_state.regime_statistics:
                        st.session_state.regime_statistics[symbol] = {
                            'regime_durations': defaultdict(list),
                            'transition_probabilities': defaultdict(lambda: defaultdict(int)),
                            'regime_returns': defaultdict(list)
                        }
                    
                    # Update regime statistics
                    transitions = st.session_state.regime_transitions[symbol]
                    if len(transitions) > 1:
                        for i in range(1, len(transitions)):
                            prev_regime = transitions[i-1]['to_regime']
                            curr_regime = transitions[i]['to_regime']
                            duration = (transitions[i]['timestamp'] - transitions[i-1]['timestamp']).total_seconds() / 3600  # hours
                            
                            st.session_state.regime_statistics[symbol]['regime_durations'][prev_regime].append(duration)
                            st.session_state.regime_statistics[symbol]['transition_probabilities'][prev_regime][curr_regime] += 1
                            
                            # Calculate returns during regime
                            regime_start = transitions[i-1]['timestamp']
                            regime_end = transitions[i]['timestamp']
                            regime_data = data[data.index.isin(pd.date_range(regime_start, regime_end))]
                            if not regime_data.empty:
                                regime_return = (regime_data['Close'].iloc[-1] - regime_data['Close'].iloc[0]) / regime_data['Close'].iloc[0]
                                st.session_state.regime_statistics[symbol]['regime_returns'][prev_regime].append(regime_return)
                    
                    # Initialize or update filter thresholds based on market regime
                    if symbol not in st.session_state.filter_thresholds:
                        st.session_state.filter_thresholds[symbol] = {
                            'volatility_threshold': 0.02,
                            'trend_threshold': 0.5,
                            'correlation_threshold': 0.7,
                            'volume_threshold': 1.5,
                            'momentum_threshold': 0.3,
                            'regime': regime
                        }
                    
                    # Update thresholds based on market regime and performance
                    current_thresholds = st.session_state.filter_thresholds[symbol]
                    
                    # Adjust thresholds based on regime statistics
                    if symbol in st.session_state.regime_statistics:
                        stats = st.session_state.regime_statistics[symbol]
                        avg_duration = np.mean(stats['regime_durations'][regime]) if stats['regime_durations'][regime] else 24
                        avg_return = np.mean(stats['regime_returns'][regime]) if stats['regime_returns'][regime] else 0
                        
                        # Adjust thresholds based on regime duration and returns
                        if avg_duration < 12:  # Short regime duration
                            current_thresholds['volatility_threshold'] *= 0.9
                            current_thresholds['trend_threshold'] *= 1.1
                        elif avg_duration > 48:  # Long regime duration
                            current_thresholds['volatility_threshold'] *= 1.1
                            current_thresholds['trend_threshold'] *= 0.9
                        
                        if avg_return > 0.02:  # Positive regime returns
                            current_thresholds['momentum_threshold'] *= 0.9
                        elif avg_return < -0.02:  # Negative regime returns
                            current_thresholds['momentum_threshold'] *= 1.1
                    
                    # Calculate market filters with adaptive thresholds
                    volatility_filter = volatility < current_thresholds['volatility_threshold']
                    trend_filter = trend_strength > current_thresholds['trend_threshold']
                    correlation_filter = correlation > current_thresholds['correlation_threshold']
                    volume_filter = volume_ratio > current_thresholds['volume_threshold']
                    momentum_filter = abs(momentum) > current_thresholds['momentum_threshold']
                    
                    # Store market filters with enhanced information
                    st.session_state.market_filters[symbol] = {
                        'volatility_filter': volatility_filter,
                        'trend_filter': trend_filter,
                        'correlation_filter': correlation_filter,
                        'volume_filter': volume_filter,
                        'momentum_filter': momentum_filter,
                        'timestamp': datetime.now(),
                        'regime': regime,
                        'regime_confidence': regime_confidence,
                        'thresholds': current_thresholds
                    }
                    
                    # Get news sentiment with caching
                    if symbol not in st.session_state.news_cache:
                        news_data = st.session_state.market_data.get_news(symbol)
                        if news_data is not None and not news_data.empty:
                            # Analyze sentiment and extract entities for each news item
                            sentiments = []
                            summaries = []
                            entities = []
                            events = []
                            
                            for _, news in news_data.iterrows():
                                # Sentiment analysis
                                sentiment = st.session_state.llm_processor.analyze_sentiment(news['title'] + ' ' + news['content'])
                                sentiments.append(sentiment)
                                
                                # Text summarization
                                summary = st.session_state.llm_processor.summarize_text(news['content'])
                                summaries.append(summary)
                                
                                # Entity extraction
                                extracted_entities = st.session_state.llm_processor.extract_entities(news['content'])
                                entities.append(extracted_entities)
                                
                                # Event detection
                                detected_events = st.session_state.llm_processor.detect_events(news['content'])
                                events.append(detected_events)
                            
                            # Store in cache
                            st.session_state.news_cache[symbol] = {
                                'data': news_data,
                                'sentiments': sentiments,
                                'summaries': summaries,
                                'entities': entities,
                                'events': events
                            }
                    
                    # Use cached data
                    news_cache = st.session_state.news_cache.get(symbol, {})
                    if news_cache:
                        news_data = news_cache['data']
                        sentiments = news_cache['sentiments']
                        summaries = news_cache['summaries']
                        entities = news_cache['entities']
                        events = news_cache['events']
                        
                        # Calculate aggregate sentiment with time decay
                        current_time = datetime.now()
                        sentiment_scores = []
                        for i, (_, news) in enumerate(news_data.iterrows()):
                            news_time = pd.to_datetime(news['date'])
                            time_diff = (current_time - news_time).total_seconds() / (24 * 3600)  # days
                            decay_factor = np.exp(-time_diff / 7)  # 7-day half-life
                            sentiment_score = sentiments[i]['score'] * (1 if sentiments[i]['label'] == 'POSITIVE' else -1)
                            sentiment_scores.append(sentiment_score * decay_factor)
                        
                        weighted_sentiment = np.mean(sentiment_scores)
                        
                        # Calculate sentiment momentum and volatility
                        sentiment_momentum = np.mean(sentiment_scores[:3]) - np.mean(sentiment_scores[3:6]) if len(sentiment_scores) >= 6 else 0
                        sentiment_volatility = np.std(sentiment_scores) if len(sentiment_scores) > 1 else 0
                        
                        # Calculate sentiment trend strength
                        sentiment_trend = np.polyfit(range(len(sentiment_scores)), sentiment_scores, 1)[0] if len(sentiment_scores) > 1 else 0
                        
                        # Calculate signal quality metrics
                        if symbol in st.session_state.signal_history:
                            historical_signals = st.session_state.signal_history[symbol]
                            if len(historical_signals) > 0:
                                # Calculate basic metrics
                                correct_signals = sum(1 for s in historical_signals if 
                                    (s['direction'] == 'BULLISH' and s['price_change'] > 0) or
                                    (s['direction'] == 'BEARISH' and s['price_change'] < 0))
                                signal_accuracy = correct_signals / len(historical_signals)
                                
                                # Calculate advanced metrics
                                returns = [s['price_change'] for s in historical_signals]
                                avg_return = np.mean(returns)
                                return_std = np.std(returns)
                                sharpe_ratio = avg_return / return_std if return_std > 0 else 0
                                
                                # Calculate win rate and profit factor
                                winning_trades = [r for r in returns if r > 0]
                                losing_trades = [r for r in returns if r < 0]
                                win_rate = len(winning_trades) / len(returns) if returns else 0
                                profit_factor = abs(sum(winning_trades) / sum(losing_trades)) if sum(losing_trades) != 0 else float('inf')
                                
                                # Calculate maximum drawdown
                                cumulative_returns = np.cumsum(returns)
                                max_drawdown = 0
                                peak = cumulative_returns[0]
                                for ret in cumulative_returns:
                                    if ret > peak:
                                        peak = ret
                                    drawdown = (peak - ret) / peak
                                    max_drawdown = max(max_drawdown, drawdown)
                                
                                # Update signal performance metrics
                                st.session_state.signal_performance[symbol] = {
                                    'accuracy': signal_accuracy,
                                    'avg_return': avg_return,
                                    'return_std': return_std,
                                    'sharpe_ratio': sharpe_ratio,
                                    'win_rate': win_rate,
                                    'profit_factor': profit_factor,
                                    'max_drawdown': max_drawdown,
                                    'total_signals': len(historical_signals)
                                }
                        
                        # Generate advanced trading signals
                        if news_cache:
                            # ... existing sentiment analysis code ...
                            
                            # Calculate signal patterns
                            if symbol not in st.session_state.signal_patterns:
                                st.session_state.signal_patterns[symbol] = {
                                    'sentiment_patterns': [],
                                    'price_patterns': [],
                                    'volume_patterns': [],
                                    'regime_patterns': []
                                }
                            
                            # Detect sentiment patterns
                            sentiment_values = [s['sentiment_score'] for s in news_cache]
                            sentiment_peaks, _ = find_peaks(sentiment_values, height=0.5)
                            sentiment_troughs, _ = find_peaks([-s for s in sentiment_values], height=0.5)
                            
                            if len(sentiment_peaks) > 0 and len(sentiment_troughs) > 0:
                                pattern = {
                                    'type': 'sentiment',
                                    'peaks': sentiment_peaks.tolist(),
                                    'troughs': sentiment_troughs.tolist(),
                                    'timestamp': datetime.now()
                                }
                                st.session_state.signal_patterns[symbol]['sentiment_patterns'].append(pattern)
                            
                            # Detect price patterns
                            price_values = data['Close'].values
                            price_peaks, _ = find_peaks(price_values, height=np.mean(price_values))
                            price_troughs, _ = find_peaks([-p for p in price_values], height=-np.mean(price_values))
                            
                            if len(price_peaks) > 0 and len(price_troughs) > 0:
                                pattern = {
                                    'type': 'price',
                                    'peaks': price_peaks.tolist(),
                                    'troughs': price_troughs.tolist(),
                                    'timestamp': datetime.now()
                                }
                                st.session_state.signal_patterns[symbol]['price_patterns'].append(pattern)
                            
                            # Detect volume patterns
                            volume_values = data['Volume'].values
                            volume_peaks, _ = find_peaks(volume_values, height=np.mean(volume_values))
                            
                            if len(volume_peaks) > 0:
                                pattern = {
                                    'type': 'volume',
                                    'peaks': volume_peaks.tolist(),
                                    'timestamp': datetime.now()
                                }
                                st.session_state.signal_patterns[symbol]['volume_patterns'].append(pattern)
                            
                            # Detect regime patterns
                            if symbol in st.session_state.regime_transitions:
                                transitions = st.session_state.regime_transitions[symbol]
                                if len(transitions) > 1:
                                    pattern = {
                                        'type': 'regime',
                                        'transitions': [t['to_regime'] for t in transitions[-5:]],
                                        'timestamp': datetime.now()
                                    }
                                    st.session_state.signal_patterns[symbol]['regime_patterns'].append(pattern)
                            
                            # Generate enhanced trading signal
                            sentiment_signal = {
                                'strength': abs(weighted_sentiment),
                                'direction': 'BULLISH' if weighted_sentiment > 0 else 'BEARISH',
                                'momentum': sentiment_momentum,
                                'volatility': sentiment_volatility,
                                'trend': sentiment_trend,
                                'confidence': min(abs(weighted_sentiment) * 2, 1.0),
                                'market_filters': {
                                    'volatility_filter': volatility_filter,
                                    'trend_filter': trend_filter,
                                    'correlation_filter': correlation_filter,
                                    'volume_filter': volume_filter,
                                    'momentum_filter': momentum_filter
                                },
                                'market_regime': regime,
                                'regime_confidence': regime_confidence,
                                'patterns': {
                                    'sentiment_patterns': len(st.session_state.signal_patterns[symbol]['sentiment_patterns']),
                                    'price_patterns': len(st.session_state.signal_patterns[symbol]['price_patterns']),
                                    'volume_patterns': len(st.session_state.signal_patterns[symbol]['volume_patterns']),
                                    'regime_patterns': len(st.session_state.signal_patterns[symbol]['regime_patterns'])
                                }
                            }
                            
                            # Calculate pattern-based confidence adjustment
                            pattern_confidence = 0
                            if sentiment_signal['patterns']['sentiment_patterns'] > 0:
                                pattern_confidence += 0.2
                            if sentiment_signal['patterns']['price_patterns'] > 0:
                                pattern_confidence += 0.3
                            if sentiment_signal['patterns']['volume_patterns'] > 0:
                                pattern_confidence += 0.2
                            if sentiment_signal['patterns']['regime_patterns'] > 0:
                                pattern_confidence += 0.3
                            
                            sentiment_signal['confidence'] *= (1 + pattern_confidence)
                            
                            # Store current signal
                            st.session_state.sentiment_signals[symbol] = sentiment_signal
                            
                            # Update signal history
                            if symbol not in st.session_state.signal_history:
                                st.session_state.signal_history[symbol] = []
                            
                            current_price = data['Close'].iloc[-1]
                            price_change = (current_price - data['Close'].iloc[-2]) / data['Close'].iloc[-2]
                            
                            st.session_state.signal_history[symbol].append({
                                'timestamp': datetime.now(),
                                'signal': sentiment_signal,
                                'price': current_price,
                                'price_change': price_change,
                                'market_filters': sentiment_signal['market_filters'],
                                'market_regime': regime,
                                'regime_confidence': regime_confidence,
                                'patterns': sentiment_signal['patterns']
                            })
                            
                            # Calculate signal performance with enhanced metrics
                            if len(st.session_state.signal_history[symbol]) > 1:
                                history = st.session_state.signal_history[symbol]
                                returns = [s['price_change'] for s in history]
                                signals = [1 if s['signal']['direction'] == 'BULLISH' else -1 for s in history]
                                
                                # Calculate signal accuracy
                                correct_signals = sum(1 for r, s in zip(returns, signals) if (r > 0 and s > 0) or (r < 0 and s < 0))
                                total_signals = len(returns)
                                accuracy = correct_signals / total_signals if total_signals > 0 else 0
                                
                                # Calculate signal returns
                                signal_returns = [r * s for r, s in zip(returns, signals)]
                                avg_return = np.mean(signal_returns) if signal_returns else 0
                                
                                # Calculate risk-adjusted metrics
                                returns_std = np.std(signal_returns) if len(signal_returns) > 1 else 0
                                sharpe_ratio = avg_return / returns_std if returns_std > 0 else 0
                                
                                # Calculate drawdown
                                cumulative_returns = np.cumsum(signal_returns)
                                running_max = np.maximum.accumulate(cumulative_returns)
                                drawdown = running_max - cumulative_returns
                                max_drawdown = np.max(drawdown) if len(drawdown) > 0 else 0
                                
                                # Calculate win rate
                                winning_trades = sum(1 for r in signal_returns if r > 0)
                                total_trades = len(signal_returns)
                                win_rate = winning_trades / total_trades if total_trades > 0 else 0
                                
                                # Calculate profit factor
                                gross_profit = sum(r for r in signal_returns if r > 0)
                                gross_loss = abs(sum(r for r in signal_returns if r < 0))
                                profit_factor = gross_profit / gross_loss if gross_loss > 0 else float('inf')
                                
                                # Store performance metrics
                                st.session_state.signal_performance[symbol] = {
                                    'accuracy': accuracy,
                                    'avg_return': avg_return,
                                    'sharpe_ratio': sharpe_ratio,
                                    'max_drawdown': max_drawdown,
                                    'win_rate': win_rate,
                                    'profit_factor': profit_factor,
                                    'total_signals': total_signals,
                                    'correct_signals': correct_signals,
                                    'timestamp': datetime.now()
                                }
                                
                                # Run backtest if we have enough data
                                if len(history) >= 20:  # Minimum data points for backtest
                                    backtest_data = pd.DataFrame({
                                        'timestamp': [s['timestamp'] for s in history],
                                        'price': [s['price'] for s in history],
                                        'signal': signals,
                                        'returns': returns,
                                        'sentiment': [s['signal']['strength'] for s in history],
                                        'confidence': [s['signal']['confidence'] for s in history],
                                        'regime': [s['market_regime'] for s in history],
                                        'regime_confidence': [s['regime_confidence'] for s in history]
                                    })
                                    
                                    # Calculate backtest metrics
                                    backtest_results = {
                                        'total_return': np.sum(signal_returns),
                                        'annualized_return': np.mean(signal_returns) * 252,  # Assuming daily data
                                        'volatility': returns_std * np.sqrt(252),
                                        'sharpe_ratio': sharpe_ratio * np.sqrt(252),
                                        'max_drawdown': max_drawdown,
                                        'win_rate': win_rate,
                                        'profit_factor': profit_factor,
                                        'avg_trade': np.mean(signal_returns),
                                        'best_trade': np.max(signal_returns),
                                        'worst_trade': np.min(signal_returns),
                                        'avg_holding_period': 1,  # Assuming daily signals
                                        'total_trades': total_trades,
                                        'winning_trades': winning_trades,
                                        'losing_trades': total_trades - winning_trades
                                    }
                                    
                                    st.session_state.backtest_results[symbol] = backtest_results
                        
                        # Display enhanced analysis
                        st.sidebar.write("\nSignal Patterns:")
                        st.sidebar.write(f"Sentiment Patterns: {sentiment_signal['patterns']['sentiment_patterns']}")
                        st.sidebar.write(f"Price Patterns: {sentiment_signal['patterns']['price_patterns']}")
                        st.sidebar.write(f"Volume Patterns: {sentiment_signal['patterns']['volume_patterns']}")
                        st.sidebar.write(f"Regime Patterns: {sentiment_signal['patterns']['regime_patterns']}")
                        
                        if symbol in st.session_state.signal_performance:
                            performance = st.session_state.signal_performance[symbol]
                            st.sidebar.write("\nSignal Performance:")
                            st.sidebar.write(f"Accuracy: {performance['accuracy']:.2%}")
                            st.sidebar.write(f"Average Return: {performance['avg_return']:.2%}")
                            st.sidebar.write(f"Sharpe Ratio: {performance['sharpe_ratio']:.2f}")
                            st.sidebar.write(f"Max Drawdown: {performance['max_drawdown']:.2%}")
                            st.sidebar.write(f"Win Rate: {performance['win_rate']:.2%}")
                            st.sidebar.write(f"Profit Factor: {performance['profit_factor']:.2f}")
                            st.sidebar.write(f"Total Signals: {performance['total_signals']}")
                            st.sidebar.write(f"Correct Signals: {performance['correct_signals']}")
                        
                        if symbol in st.session_state.backtest_results:
                            backtest = st.session_state.backtest_results[symbol]
                            st.sidebar.write("\nBacktest Results:")
                            st.sidebar.write(f"Total Return: {backtest['total_return']:.2%}")
                            st.sidebar.write(f"Annualized Return: {backtest['annualized_return']:.2%}")
                            st.sidebar.write(f"Volatility: {backtest['volatility']:.2%}")
                            st.sidebar.write(f"Sharpe Ratio: {backtest['sharpe_ratio']:.2f}")
                            st.sidebar.write(f"Max Drawdown: {backtest['max_drawdown']:.2%}")
                            st.sidebar.write(f"Win Rate: {backtest['win_rate']:.2%}")
                            st.sidebar.write(f"Profit Factor: {backtest['profit_factor']:.2f}")
                            st.sidebar.write(f"Average Trade: {backtest['avg_trade']:.2%}")
                            st.sidebar.write(f"Best Trade: {backtest['best_trade']:.2%}")
                            st.sidebar.write(f"Worst Trade: {backtest['worst_trade']:.2%}")
                            st.sidebar.write(f"Total Trades: {backtest['total_trades']}")
                            st.sidebar.write(f"Winning Trades: {backtest['winning_trades']}")
                            st.sidebar.write(f"Losing Trades: {backtest['losing_trades']}")
                        
                        # Calculate portfolio optimization metrics
                        if symbol not in st.session_state.portfolio_weights:
                            st.session_state.portfolio_weights[symbol] = {
                                'current_weight': 0.0,
                                'target_weight': 0.0,
                                'max_weight': 0.2,  # Maximum 20% allocation per position
                                'min_weight': 0.0,
                                'last_update': datetime.now()
                            }
                        
                        # Calculate risk metrics
                        if symbol not in st.session_state.risk_metrics:
                            st.session_state.risk_metrics[symbol] = {
                                'var_95': 0.0,  # Value at Risk at 95% confidence
                                'var_99': 0.0,  # Value at Risk at 99% confidence
                                'expected_shortfall': 0.0,
                                'beta': 0.0,
                                'correlation': 0.0,
                                'volatility': 0.0,
                                'sharpe_ratio': 0.0,
                                'sortino_ratio': 0.0,
                                'max_drawdown': 0.0,
                                'last_update': datetime.now()
                            }
                        
                        # Calculate returns
                        returns = data['Close'].pct_change().dropna()
                        market_returns = market_data_df['Close'].pct_change().dropna()
                        
                        # Calculate risk metrics
                        volatility = returns.std() * np.sqrt(252)  # Annualized volatility
                        beta = returns.cov(market_returns) / market_returns.var()
                        correlation = returns.correl(market_returns)
                        
                        # Calculate Value at Risk
                        var_95 = norm.ppf(0.05, returns.mean(), returns.std())
                        var_99 = norm.ppf(0.01, returns.mean(), returns.std())
                        
                        # Calculate Expected Shortfall (CVaR)
                        expected_shortfall = -returns[returns <= var_95].mean()
                        
                        # Calculate Sharpe Ratio
                        risk_free_rate = 0.02  # Assuming 2% risk-free rate
                        excess_returns = returns - risk_free_rate/252
                        sharpe_ratio = np.sqrt(252) * excess_returns.mean() / returns.std()
                        
                        # Calculate Sortino Ratio
                        downside_returns = returns[returns < 0]
                        sortino_ratio = np.sqrt(252) * excess_returns.mean() / downside_returns.std()
                        
                        # Calculate Maximum Drawdown
                        cumulative_returns = (1 + returns).cumprod()
                        running_max = cumulative_returns.expanding().max()
                        drawdown = (cumulative_returns / running_max) - 1
                        max_drawdown = drawdown.min()
                        
                        # Update risk metrics
                        st.session_state.risk_metrics[symbol].update({
                            'var_95': var_95,
                            'var_99': var_99,
                            'expected_shortfall': expected_shortfall,
                            'beta': beta,
                            'correlation': correlation,
                            'volatility': volatility,
                            'sharpe_ratio': sharpe_ratio,
                            'sortino_ratio': sortino_ratio,
                            'max_drawdown': max_drawdown,
                            'last_update': datetime.now()
                        })
                        
                        # Portfolio optimization
                        if symbol in st.session_state.signal_performance:
                            performance = st.session_state.signal_performance[symbol]
                            
                            # Calculate target weight based on signal performance and risk metrics
                            risk_metrics = st.session_state.risk_metrics[symbol]
                            
                            # Score components
                            signal_score = (
                                performance['accuracy'] * 
                                (1 + performance['avg_return']) * 
                                (1 - performance['max_drawdown']) * 
                                performance['sharpe_ratio']
                            )
                            
                            risk_score = (
                                (1 - abs(risk_metrics['beta'])) *  # Lower beta is better
                                (1 - risk_metrics['volatility']) *  # Lower volatility is better
                                (1 + risk_metrics['sharpe_ratio']) *  # Higher Sharpe is better
                                (1 + risk_metrics['sortino_ratio'])  # Higher Sortino is better
                            )
                            
                            # Calculate target weight
                            base_weight = min(signal_score * risk_score, 1.0)
                            target_weight = base_weight * st.session_state.portfolio_weights[symbol]['max_weight']
                            
                            # Apply constraints
                            target_weight = max(
                                min(target_weight, st.session_state.portfolio_weights[symbol]['max_weight']),
                                st.session_state.portfolio_weights[symbol]['min_weight']
                            )
                            
                            # Update portfolio weights
                            st.session_state.portfolio_weights[symbol].update({
                                'target_weight': target_weight,
                                'last_update': datetime.now()
                            })
                            
                            # Run portfolio optimization if we have enough data
                            if len(st.session_state.signal_history[symbol]) >= 20:
                                # Prepare optimization data
                                history = st.session_state.signal_history[symbol]
                                returns_data = pd.DataFrame({
                                    'timestamp': [s['timestamp'] for s in history],
                                    'returns': [s['price_change'] for s in history],
                                    'sentiment': [s['signal']['strength'] for s in history],
                                    'confidence': [s['signal']['confidence'] for s in history],
                                    'regime': [s['market_regime'] for s in history]
                                })
                                
                                # Calculate regime-specific returns
                                regime_returns = returns_data.groupby('regime')['returns'].agg(['mean', 'std']).to_dict()
                                
                                # Calculate regime transition probabilities
                                if symbol in st.session_state.regime_statistics:
                                    stats = st.session_state.regime_statistics[symbol]
                                    transition_probs = stats['transition_probabilities']
                                    
                                    # Convert to probability matrix
                                    regimes = list(transition_probs.keys())
                                    prob_matrix = np.zeros((len(regimes), len(regimes)))
                                    
                                    for i, from_regime in enumerate(regimes):
                                        total = sum(transition_probs[from_regime].values())
                                        if total > 0:
                                            for j, to_regime in enumerate(regimes):
                                                prob_matrix[i, j] = transition_probs[from_regime].get(to_regime, 0) / total
                            
                                # Run optimization
                                def objective(weights):
                                    # Calculate portfolio metrics
                                    portfolio_return = np.sum(weights * returns.mean())
                                    portfolio_vol = np.sqrt(weights.T @ returns.cov() @ weights)
                                    sharpe = (portfolio_return - risk_free_rate) / portfolio_vol
                                    return -sharpe  # Minimize negative Sharpe ratio
                                
                                # Constraints
                                constraints = [
                                    {'type': 'eq', 'fun': lambda x: np.sum(x) - 1},  # Weights sum to 1
                                    {'type': 'ineq', 'fun': lambda x: x - st.session_state.portfolio_weights[symbol]['min_weight']},  # Minimum weight
                                    {'type': 'ineq', 'fun': lambda x: st.session_state.portfolio_weights[symbol]['max_weight'] - x}  # Maximum weight
                                ]
                                
                                # Initial guess
                                initial_weights = np.array([0.5])
                                
                                # Run optimization
                                result = minimize(
                                    objective,
                                    initial_weights,
                                    method='SLSQP',
                                    constraints=constraints,
                                    bounds=[(0, 1)]
                                )
                                
                                if result.success:
                                    optimal_weight = result.x[0]
                                    
                                    # Store optimization results
                                    st.session_state.optimization_results[symbol] = {
                                        'optimal_weight': optimal_weight,
                                        'expected_return': -result.fun,
                                        'constraints_satisfied': all(c['fun'](optimal_weight) >= 0 for c in constraints),
                                        'optimization_status': result.message,
                                        'last_update': datetime.now()
                                    }
                                    
                                    # Update target weight with optimization result
                                    st.session_state.portfolio_weights[symbol]['target_weight'] = optimal_weight
                    
                        # Display enhanced analysis
                        st.sidebar.write("\nRisk Metrics:")
                        risk_metrics = st.session_state.risk_metrics[symbol]
                        st.sidebar.write(f"Value at Risk (95%): {risk_metrics['var_95']:.2%}")
                        st.sidebar.write(f"Value at Risk (99%): {risk_metrics['var_99']:.2%}")
                        st.sidebar.write(f"Expected Shortfall: {risk_metrics['expected_shortfall']:.2%}")
                        st.sidebar.write(f"Beta: {risk_metrics['beta']:.2f}")
                        st.sidebar.write(f"Correlation: {risk_metrics['correlation']:.2f}")
                        st.sidebar.write(f"Volatility: {risk_metrics['volatility']:.2%}")
                        st.sidebar.write(f"Sharpe Ratio: {risk_metrics['sharpe_ratio']:.2f}")
                        st.sidebar.write(f"Sortino Ratio: {risk_metrics['sortino_ratio']:.2f}")
                        st.sidebar.write(f"Max Drawdown: {risk_metrics['max_drawdown']:.2%}")
                        
                        st.sidebar.write("\nPortfolio Allocation:")
                        portfolio_weights = st.session_state.portfolio_weights[symbol]
                        st.sidebar.write(f"Current Weight: {portfolio_weights['current_weight']:.2%}")
                        st.sidebar.write(f"Target Weight: {portfolio_weights['target_weight']:.2%}")
                        st.sidebar.write(f"Max Weight: {portfolio_weights['max_weight']:.2%}")
                        
                        if symbol in st.session_state.optimization_results:
                            optimization = st.session_state.optimization_results[symbol]
                            st.sidebar.write("\nOptimization Results:")
                            st.sidebar.write(f"Optimal Weight: {optimization['optimal_weight']:.2%}")
                            st.sidebar.write(f"Expected Return: {optimization['expected_return']:.2%}")
                            st.sidebar.write(f"Constraints Satisfied: {optimization['constraints_satisfied']}")
                            st.sidebar.write(f"Optimization Status: {optimization['optimization_status']}")
                        
                        # Initialize execution strategy
                        if symbol not in st.session_state.execution_strategies:
                            st.session_state.execution_strategies[symbol] = {
                                'strategy_type': 'ADAPTIVE',
                                'parameters': {
                                    'urgency': 0.5,  # 0-1 scale, higher means more aggressive
                                    'participation_rate': 0.1,  # Target participation in volume
                                    'max_slippage': 0.001,  # Maximum allowed slippage
                                    'min_spread': 0.0005,  # Minimum spread to consider
                                    'time_horizon': 3600,  # Execution time horizon in seconds
                                    'price_improvement': 0.0002,  # Target price improvement
                                    'risk_aversion': 0.5,  # Risk aversion parameter
                                    'market_impact': 0.0001,  # Estimated market impact
                                    'last_update': datetime.now()
                                }
                            }
                        
                        # Calculate execution parameters based on market conditions
                        if symbol in st.session_state.risk_metrics:
                            risk_metrics = st.session_state.risk_metrics[symbol]
                            market_conditions = st.session_state.market_analyzer.analyze_market_conditions(data, market_data_df)
                            
                            # Adjust execution parameters based on volatility
                            volatility = risk_metrics['volatility']
                            if volatility > 0.3:  # High volatility
                                st.session_state.execution_strategies[symbol]['parameters']['urgency'] = 0.8
                                st.session_state.execution_strategies[symbol]['parameters']['participation_rate'] = 0.05
                                st.session_state.execution_strategies[symbol]['parameters']['max_slippage'] = 0.002
                            elif volatility > 0.2:  # Medium volatility
                                st.session_state.execution_strategies[symbol]['parameters']['urgency'] = 0.6
                                st.session_state.execution_strategies[symbol]['parameters']['participation_rate'] = 0.1
                                st.session_state.execution_strategies[symbol]['parameters']['max_slippage'] = 0.001
                            else:  # Low volatility
                                st.session_state.execution_strategies[symbol]['parameters']['urgency'] = 0.4
                                st.session_state.execution_strategies[symbol]['parameters']['participation_rate'] = 0.15
                                st.session_state.execution_strategies[symbol]['parameters']['max_slippage'] = 0.0005
                            
                            # Adjust parameters based on market regime
                            if symbol in st.session_state.market_regime:
                                regime = st.session_state.market_regime[symbol]['regime']
                                if regime == 'HIGH_VOLATILITY':
                                    st.session_state.execution_strategies[symbol]['strategy_type'] = 'AGGRESSIVE'
                                    st.session_state.execution_strategies[symbol]['parameters']['time_horizon'] = 1800  # 30 minutes
                                elif regime == 'TRENDING':
                                    st.session_state.execution_strategies[symbol]['strategy_type'] = 'PASSIVE'
                                    st.session_state.execution_strategies[symbol]['parameters']['time_horizon'] = 7200  # 2 hours
                                else:
                                    st.session_state.execution_strategies[symbol]['strategy_type'] = 'ADAPTIVE'
                                    st.session_state.execution_strategies[symbol]['parameters']['time_horizon'] = 3600  # 1 hour
                            
                            # Calculate market impact
                            volume = data['Volume'].iloc[-1]
                            avg_volume = data['Volume'].rolling(20).mean().iloc[-1]
                            volume_ratio = volume / avg_volume
                            
                            # Adjust market impact based on volume
                            if volume_ratio > 2:
                                st.session_state.execution_strategies[symbol]['parameters']['market_impact'] *= 0.5
                            elif volume_ratio < 0.5:
                                st.session_state.execution_strategies[symbol]['parameters']['market_impact'] *= 2
                            
                            # Update execution parameters
                            st.session_state.execution_strategies[symbol]['parameters']['last_update'] = datetime.now()
                        
                        # Generate execution orders based on portfolio weights
                        if symbol in st.session_state.portfolio_weights:
                            portfolio_weights = st.session_state.portfolio_weights[symbol]
                            current_weight = portfolio_weights['current_weight']
                            target_weight = portfolio_weights['target_weight']
                            
                            if abs(target_weight - current_weight) > 0.001:  # Minimum threshold for rebalancing
                                # Calculate order size
                                current_price = data['Close'].iloc[-1]
                                order_size = abs(target_weight - current_weight)
                                
                                # Determine order type based on execution strategy
                                execution_strategy = st.session_state.execution_strategies[symbol]
                                if execution_strategy['strategy_type'] == 'AGGRESSIVE':
                                    order_type = OrderType.MARKET
                                elif execution_strategy['strategy_type'] == 'PASSIVE':
                                    order_type = OrderType.LIMIT
                                else:  # ADAPTIVE
                                    if abs(target_weight - current_weight) > 0.05:  # Large position change
                                        order_type = OrderType.TWAP
                                    else:
                                        order_type = OrderType.LIMIT
                                
                                # Create order
                                order_id = str(uuid.uuid4())
                                order = {
                                    'id': order_id,
                                    'symbol': symbol,
                                    'type': order_type.value,
                                    'side': 'BUY' if target_weight > current_weight else 'SELL',
                                    'size': order_size,
                                    'price': current_price,
                                    'status': OrderStatus.PENDING.value,
                                    'created_at': datetime.now(),
                                    'execution_strategy': execution_strategy,
                                    'parameters': {
                                        'urgency': execution_strategy['parameters']['urgency'],
                                        'participation_rate': execution_strategy['parameters']['participation_rate'],
                                        'max_slippage': execution_strategy['parameters']['max_slippage'],
                                        'time_horizon': execution_strategy['parameters']['time_horizon']
                                    }
                                }
                                
                                # Store order
                                st.session_state.orders[order_id] = order
                                
                                # Initialize order history
                                if symbol not in st.session_state.order_history:
                                    st.session_state.order_history[symbol] = []
                                
                                st.session_state.order_history[symbol].append({
                                    'order_id': order_id,
                                    'type': order_type.value,
                                    'side': order['side'],
                                    'size': order_size,
                                    'price': current_price,
                                    'status': OrderStatus.PENDING.value,
                                    'created_at': datetime.now(),
                                    'execution_strategy': execution_strategy['strategy_type']
                                })
                        
                        # Display enhanced analysis
                        st.sidebar.write("\nExecution Strategy:")
                        execution_strategy = st.session_state.execution_strategies[symbol]
                        st.sidebar.write(f"Strategy Type: {execution_strategy['strategy_type']}")
                        st.sidebar.write(f"Urgency: {execution_strategy['parameters']['urgency']:.2f}")
                        st.sidebar.write(f"Participation Rate: {execution_strategy['parameters']['participation_rate']:.2%}")
                        st.sidebar.write(f"Max Slippage: {execution_strategy['parameters']['max_slippage']:.2%}")
                        st.sidebar.write(f"Time Horizon: {execution_strategy['parameters']['time_horizon']/3600:.1f} hours")
                        st.sidebar.write(f"Price Improvement Target: {execution_strategy['parameters']['price_improvement']:.2%}")
                        st.sidebar.write(f"Market Impact: {execution_strategy['parameters']['market_impact']:.2%}")
                        
                        if symbol in st.session_state.order_history:
                            st.sidebar.write("\nRecent Orders:")
                            for order in st.session_state.order_history[symbol][-5:]:  # Show last 5 orders
                                st.sidebar.write(f"Order ID: {order['order_id'][:8]}...")
                                st.sidebar.write(f"Type: {order['type']}")
                                st.sidebar.write(f"Side: {order['side']}")
                                st.sidebar.write(f"Size: {order['size']:.2%}")
                                st.sidebar.write(f"Price: ${order['price']:.2f}")
                                st.sidebar.write(f"Status: {order['status']}")
                                st.sidebar.write(f"Strategy: {order['execution_strategy']}")
                                st.sidebar.write("---")
                        
                        # Create enhanced visualization
                        fig = make_subplots(rows=13, cols=1,
                                          shared_xaxes=True,
                                          vertical_spacing=0.05,
                                          subplot_titles=('Price and Trend', 'Volatility', 'Correlation', 
                                                        'News Sentiment', 'Entity Mentions', 'Trading Signals',
                                                        'Signal Performance', 'Market Filters', 'Market Regime',
                                                        'Regime Transitions', 'Signal Patterns', 'Risk Metrics',
                                                        'Execution Metrics'),
                                          row_heights=[0.2, 0.15, 0.15, 0.15, 0.15, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1])
                        
                        # Price and Trend plot
                        fig.add_trace(go.Scatter(x=data.index, y=data['Close'],
                                               name='Price', line=dict(color='blue')),
                                    row=1, col=1)
                        fig.add_trace(go.Scatter(x=data.index, y=data['Close'].rolling(20).mean(),
                                               name='20-day MA', line=dict(color='orange')),
                                    row=1, col=1)
                        fig.add_trace(go.Scatter(x=data.index, y=data['Close'].rolling(50).mean(),
                                               name='50-day MA', line=dict(color='red')),
                                    row=1, col=1)
                        
                        # Volatility plot
                        returns = data['Close'].pct_change()
                        volatility = returns.rolling(window=20).std() * np.sqrt(252)
                        fig.add_trace(go.Scatter(x=data.index, y=volatility,
                                               name='Volatility', line=dict(color='green')),
                                    row=2, col=1)
                        
                        # Correlation plot
                        market_returns = market_data_df['Close'].pct_change()
                        correlation = returns.rolling(window=20).corr(market_returns)
                        fig.add_trace(go.Scatter(x=data.index, y=correlation,
                                               name='Correlation', line=dict(color='purple')),
                                    row=3, col=1)
                        
                        # Enhanced news sentiment plot
                        if news_cache:
                            sentiment_dates = pd.to_datetime(news_data['date'])
                            fig.add_trace(go.Scatter(x=sentiment_dates, y=sentiment_scores,
                                                   name='Weighted Sentiment', line=dict(color='orange')),
                                        row=4, col=1)
                            
                            # Entity mentions plot
                            entity_dates = []
                            entity_values = []
                            for i, (_, news) in enumerate(news_data.iterrows()):
                                entity_dates.append(pd.to_datetime(news['date']))
                                entity_values.append(len(entities[i]))
                            
                            fig.add_trace(go.Scatter(x=entity_dates, y=entity_values,
                                                   name='Entity Mentions', line=dict(color='green')),
                                        row=5, col=1)
                            
                            # Trading signals plot
                            signal_dates = sentiment_dates[:len(sentiment_scores)]
                            signal_values = [1 if s > 0 else -1 for s in sentiment_scores]
                            fig.add_trace(go.Scatter(x=signal_dates, y=signal_values,
                                                   name='Trading Signals', line=dict(color='purple')),
                                        row=6, col=1)
                            
                            # Signal performance plot
                            if symbol in st.session_state.signal_history:
                                history = st.session_state.signal_history[symbol]
                                perf_dates = [s['timestamp'] for s in history]
                                perf_values = [s['price_change'] for s in history]
                                fig.add_trace(go.Scatter(x=perf_dates, y=perf_values,
                                                       name='Signal Returns', line=dict(color='blue')),
                                            row=7, col=1)
                                
                                # Market filters plot
                                filter_values = [
                                    sum([
                                        s['market_filters']['volatility_filter'],
                                        s['market_filters']['trend_filter'],
                                        s['market_filters']['correlation_filter'],
                                        s['market_filters']['volume_filter'],
                                        s['market_filters']['momentum_filter']
                                    ]) / 5.0 for s in history
                                ]
                                fig.add_trace(go.Scatter(x=perf_dates, y=filter_values,
                                                       name='Filter Score', line=dict(color='green')),
                                            row=8, col=1)
                                
                                # Market regime plot
                                regime_values = [
                                    1 if s.get('market_regime') == 'TRENDING' else
                                    0.8 if s.get('market_regime') == 'CORRELATED' else
                                    0.6 if s.get('market_regime') == 'NORMAL' else
                                    0.4 for s in history
                                ]
                                fig.add_trace(go.Scatter(x=perf_dates, y=regime_values,
                                                       name='Market Regime', line=dict(color='purple')),
                                            row=9, col=1)
                                
                                # Regime transitions plot
                                if symbol in st.session_state.regime_transitions:
                                    transitions = st.session_state.regime_transitions[symbol]
                                    trans_dates = [t['timestamp'] for t in transitions]
                                    trans_values = [t['confidence'] for t in transitions]
                                    fig.add_trace(go.Scatter(x=trans_dates, y=trans_values,
                                                           name='Regime Transitions', line=dict(color='orange')),
                                                row=10, col=1)
                        
                        # Signal patterns plot
                        if symbol in st.session_state.signal_patterns:
                            patterns = st.session_state.signal_patterns[symbol]
                            if patterns['sentiment_patterns']:
                                last_pattern = patterns['sentiment_patterns'][-1]
                                fig.add_trace(go.Scatter(x=data.index[last_pattern['peaks']], 
                                                       y=data['Close'].iloc[last_pattern['peaks']],
                                                       mode='markers', name='Sentiment Peaks',
                                                       marker=dict(color='green', size=10)),
                                            row=1, col=1)
                            
                            if patterns['price_patterns']:
                                last_pattern = patterns['price_patterns'][-1]
                                fig.add_trace(go.Scatter(x=data.index[last_pattern['peaks']], 
                                                       y=data['Close'].iloc[last_pattern['peaks']],
                                                       mode='markers', name='Price Peaks',
                                                       marker=dict(color='blue', size=10)),
                                            row=1, col=1)
                            
                            if patterns['volume_patterns']:
                                last_pattern = patterns['volume_patterns'][-1]
                                fig.add_trace(go.Scatter(x=data.index[last_pattern['peaks']], 
                                                       y=data['Volume'].iloc[last_pattern['peaks']],
                                                       mode='markers', name='Volume Peaks',
                                                       marker=dict(color='orange', size=10)),
                                            row=2, col=1)
                        
                        # Risk metrics plot
                        if symbol in st.session_state.risk_metrics:
                            risk_metrics = st.session_state.risk_metrics[symbol]
                            metrics = ['var_95', 'var_99', 'expected_shortfall', 'volatility']
                            values = [abs(risk_metrics[m]) for m in metrics]
                            fig.add_trace(go.Bar(x=metrics, y=values, name='Risk Metrics'),
                                        row=11, col=1)
                        
                        # Execution metrics plot
                        if symbol in st.session_state.order_history:
                            orders = st.session_state.order_history[symbol]
                            if orders:
                                order_dates = [o['created_at'] for o in orders]
                                order_sizes = [o['size'] for o in orders]
                                order_prices = [o['price'] for o in orders]
                                
                                fig.add_trace(go.Scatter(x=order_dates, y=order_sizes,
                                                       name='Order Sizes', line=dict(color='purple')),
                                            row=12, col=1)
                                
                                fig.add_trace(go.Scatter(x=order_dates, y=order_prices,
                                                       name='Order Prices', line=dict(color='orange')),
                                            row=12, col=1)
                        
                        fig.update_layout(height=2800, title_text=f'Enhanced Market Analysis for {symbol}')
                        st.plotly_chart(fig)
            
            elif action['type'] == 'train_model':
                symbol = action.get('symbol', 'AAPL')
                epochs = action.get('epochs', 100)
                batch_size = action.get('batch_size', 32)
                
                # Get historical data for training
                data = st.session_state.market_data.get_historical_data(symbol, window=100)  # Get more data for training
                if data is not None and not data.empty:
                    # Prepare data for training
                    feature_data = data[['open', 'high', 'low', 'close', 'volume']].copy()
                    target_data = data['close'].values
                    
                    # Train the model
                    training_metrics = st.session_state.lstm_model.fit(
                        feature_data, target_data,
                        epochs=epochs,
                        batch_size=batch_size
                    )
                    
                    # Store metrics
                    st.session_state.model_metrics[symbol] = training_metrics
                    
                    # Display training progress
                    st.sidebar.write(f"Training completed for {symbol}")
                    st.sidebar.write("Training Metrics:")
                    for metric, value in training_metrics.items():
                        st.sidebar.write(f"{metric}: {value:.4f}")
                    
                    # Plot training metrics
                    fig = make_subplots(rows=2, cols=1, subplot_titles=('Training Loss', 'Validation Loss'))
                    fig.add_trace(go.Scatter(y=training_metrics['train_loss'], name='Training Loss'), row=1, col=1)
                    fig.add_trace(go.Scatter(y=training_metrics['val_loss'], name='Validation Loss'), row=1, col=1)
                    fig.update_layout(height=600, title_text="Model Training Progress")
                    st.plotly_chart(fig)
            
            elif action['type'] == 'predict_price':
                symbol = action.get('symbol', 'AAPL')
                window = action.get('window', 10)
                
                # Get historical data
                data = st.session_state.market_data.get_historical_data(symbol, window)
                if data is not None and not data.empty:
                    # Prepare data for prediction
                    feature_data = data[['open', 'high', 'low', 'close', 'volume']].copy()
                    
                    # Make prediction
                    prediction = st.session_state.lstm_model.predict(feature_data)
                    
                    # Calculate metrics if we have actual values
                    if symbol in st.session_state.model_metrics:
                        metrics = st.session_state.model_metrics[symbol]
                        st.sidebar.write("Model Performance Metrics:")
                        st.sidebar.write(f"RMSE: {metrics['rmse']:.4f}")
                        st.sidebar.write(f"MAE: {metrics['mae']:.4f}")
                        st.sidebar.write(f"R² Score: {metrics['r2']:.4f}")
                    
                    # Display prediction
                    st.sidebar.write(f"LSTM Prediction for {symbol}:")
                    st.sidebar.write(f"Predicted price: ${prediction[-1]:.2f}")
                    
                    # Create advanced visualization
                    fig = make_subplots(rows=2, cols=1, 
                                      shared_xaxes=True,
                                      vertical_spacing=0.03,
                                      subplot_titles=('Price Prediction', 'Prediction Error'),
                                      row_heights=[0.7, 0.3])
                    
                    # Price plot
                    fig.add_trace(go.Scatter(x=data.index, y=data['close'], 
                                           name='Actual', line=dict(color='blue')),
                                row=1, col=1)
                    fig.add_trace(go.Scatter(x=data.index, y=prediction, 
                                           name='Predicted', line=dict(color='red')),
                                row=1, col=1)
                    
                    # Error plot
                    error = data['close'].values - prediction
                    fig.add_trace(go.Scatter(x=data.index, y=error,
                                           name='Prediction Error',
                                           line=dict(color='green')),
                                row=2, col=1)
                    
                    # Add confidence intervals
                    std_error = np.std(error)
                    fig.add_trace(go.Scatter(x=data.index, y=prediction + 2*std_error,
                                           name='Upper Bound',
                                           line=dict(color='gray', dash='dash')),
                                row=1, col=1)
                    fig.add_trace(go.Scatter(x=data.index, y=prediction - 2*std_error,
                                           name='Lower Bound',
                                           line=dict(color='gray', dash='dash')),
                                row=1, col=1)
                    
                    fig.update_layout(height=800, title_text=f'LSTM Price Prediction Analysis for {symbol}')
                    st.plotly_chart(fig)
            
            elif action['type'] == 'evaluate_model':
                symbol = action.get('symbol', 'AAPL')
                if symbol in st.session_state.model_metrics:
                    metrics = st.session_state.model_metrics[symbol]
                    
                    # Create metrics visualization
                    fig = go.Figure()
                    metrics_to_plot = ['rmse', 'mae', 'r2']
                    fig.add_trace(go.Bar(x=metrics_to_plot,
                                       y=[metrics[m] for m in metrics_to_plot],
                                       text=[f'{metrics[m]:.4f}' for m in metrics_to_plot],
                                       textposition='auto'))
                    
                    fig.update_layout(title=f'Model Performance Metrics for {symbol}',
                                    yaxis_title='Value',
                                    showlegend=False)
                    st.plotly_chart(fig)
                    
                    # Display detailed metrics
                    st.sidebar.write("Detailed Model Evaluation:")
                    for metric, value in metrics.items():
                        st.sidebar.write(f"{metric}: {value:.4f}")

# Main navigation
page = st.sidebar.radio(
    "Navigation",
    ["Dashboard", "Trading", "Backtesting", "Risk Management", 
     "Portfolio", "Analysis", "Optimization", "ML Models", "Settings"]
)

# Main content
if page == "Dashboard":
    st.title("Trading Dashboard")
    
    # Portfolio Overview
    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Total Value", "$1,234,567", "+2.3%")
    with col2:
        st.metric("Daily P/L", "$12,345", "-1.2%")
    with col3:
        st.metric("Open Positions", "15", "+2")
    with col4:
        st.metric("Risk Score", "0.65", "-0.05")
    
    # Performance Chart
    st.subheader("Portfolio Performance")
    dates = pd.date_range(start='2024-01-01', end='2024-01-31', freq='D')
    portfolio_values = pd.Series(range(len(dates))) * 1000 + 1000000
    benchmark_values = pd.Series(range(len(dates))) * 800 + 1000000
    df_perf = pd.DataFrame({
        'Date': dates,
        'Portfolio': portfolio_values,
        'Benchmark': benchmark_values
    })
    fig = px.line(df_perf, x='Date', y=['Portfolio', 'Benchmark'],
                  title='Portfolio vs Benchmark')
    st.plotly_chart(fig, use_container_width=True)
    
    # Recent Trades
    st.subheader("Recent Trades")
    trades_data = {
        'Time': pd.date_range(start='2024-01-01', periods=5, freq='H'),
        'Symbol': ['AAPL', 'GOOGL', 'MSFT', 'AMZN', 'META'],
        'Type': ['Buy', 'Sell', 'Buy', 'Sell', 'Buy'],
        'Quantity': [100, 50, 75, 25, 150],
        'Price': [180.5, 140.2, 380.0, 150.75, 320.25],
        'P/L': [1200, -500, 800, -300, 1500]
    }
    st.dataframe(pd.DataFrame(trades_data))

elif page == "Trading":
    st.title("Trading")
    
    # Strategy Selection
    st.subheader("Active Strategies")
    strategies = ["Momentum", "Mean Reversion", "ML-Based", "Custom"]
    selected_strategies = st.multiselect("Select Strategies", strategies)
    
    # Trading Parameters
    col1, col2 = st.columns(2)
    with col1:
        st.number_input("Position Size (%)", 1, 100, 5)
        st.number_input("Stop Loss (%)", 1, 20, 5)
    with col2:
        st.number_input("Take Profit (%)", 1, 50, 10)
        st.selectbox("Timeframe", ["1m", "5m", "15m", "1h", "4h", "1d"])
    
    # Market Scanner
    st.subheader("Market Scanner")
    scanner_options = ["Top Gainers", "Top Losers", "Volume Leaders", "Technical Signals"]
    selected_scanner = st.selectbox("Scanner Type", scanner_options)
    
    # Display scanner results
    scanner_data = {
        'Symbol': ['AAPL', 'GOOGL', 'MSFT', 'AMZN', 'META'],
        'Price': [180.5, 140.2, 380.0, 150.75, 320.25],
        'Change %': [2.3, -1.5, 3.2, -0.8, 1.9],
        'Volume': [1000000, 800000, 1200000, 900000, 1100000],
        'Signal': ['Buy', 'Sell', 'Buy', 'Hold', 'Buy']
    }
    st.dataframe(pd.DataFrame(scanner_data))

elif page == "Backtesting":
    st.title("Backtesting")
    
    # Backtest Configuration
    col1, col2 = st.columns(2)
    with col1:
        st.date_input("Start Date", datetime(2023, 1, 1))
        st.selectbox("Strategy", ["Momentum", "Mean Reversion", "ML-Based"])
    with col2:
        st.date_input("End Date", datetime(2024, 1, 1))
        st.multiselect("Assets", ["AAPL", "GOOGL", "MSFT", "AMZN", "META"])
    
    # Run Backtest
    if st.button("Run Backtest"):
        st.subheader("Backtest Results")
        
        # Performance Metrics
        metrics = {
            'Metric': ['Total Return', 'Sharpe Ratio', 'Max Drawdown', 'Win Rate'],
            'Value': ['23.5%', '1.8', '12.3%', '65%']
        }
        st.dataframe(pd.DataFrame(metrics))
        
        # Performance Chart
        dates = pd.date_range(start='2023-01-01', end='2024-01-01', freq='M')
        strategy_values = pd.Series(range(len(dates))) * 100 + 10000
        benchmark_values = pd.Series(range(len(dates))) * 80 + 10000
        df_bt = pd.DataFrame({
            'Date': dates,
            'Strategy': strategy_values,
            'Benchmark': benchmark_values
        })
        fig = px.line(df_bt, x='Date', y=['Strategy', 'Benchmark'],
                      title='Backtest Performance')
        st.plotly_chart(fig, use_container_width=True)

elif page == "Risk Management":
    st.title("Risk Management")
    
    # Risk Metrics
    col1, col2, col3 = st.columns(3)
    with col1:
        st.metric("Portfolio VaR", "2.3%", "-0.5%")
    with col2:
        st.metric("Beta", "1.2", "+0.1")
    with col3:
        st.metric("Correlation", "0.75", "-0.05")
    
    # Risk Analysis
    st.subheader("Risk Analysis")
    risk_factors = ["Market Risk", "Liquidity Risk", "Volatility Risk", "Credit Risk"]
    risk_scores = [0.65, 0.45, 0.75, 0.35]
    
    fig = go.Figure(data=[
        go.Bar(x=risk_factors, y=risk_scores)
    ])
    fig.update_layout(title="Risk Factor Analysis")
    st.plotly_chart(fig, use_container_width=True)
    
    # Position Limits
    st.subheader("Position Limits")
    limits_data = {
        'Asset': ['AAPL', 'GOOGL', 'MSFT', 'AMZN', 'META'],
        'Current %': [15, 12, 10, 8, 7],
        'Max %': [20, 15, 15, 10, 10],
        'Risk Score': [0.65, 0.75, 0.55, 0.85, 0.70]
    }
    st.dataframe(pd.DataFrame(limits_data))

elif page == "Portfolio":
    st.title("Portfolio Management")
    
    # Portfolio Composition
    st.subheader("Portfolio Composition")
    assets = ['AAPL', 'GOOGL', 'MSFT', 'AMZN', 'META', 'Others']
    weights = [25, 20, 15, 15, 15, 10]
    
    fig = px.pie(values=weights, names=assets, title="Portfolio Allocation")
    st.plotly_chart(fig, use_container_width=True)
    
    # Holdings
    st.subheader("Current Holdings")
    holdings_data = {
        'Asset': ['AAPL', 'GOOGL', 'MSFT', 'AMZN', 'META'],
        'Quantity': [100, 50, 75, 25, 150],
        'Avg Price': [150.5, 2800.2, 280.0, 3300.75, 180.25],
        'Current Price': [180.5, 2900.2, 300.0, 3400.75, 200.25],
        'P/L': [3000, 5000, 1500, 2500, 3000],
        'P/L %': [20, 3.5, 7.1, 3.0, 11.1]
    }
    st.dataframe(pd.DataFrame(holdings_data))

elif page == "Analysis":
    st.title("Market Analysis")
    
    # Technical Analysis
    st.subheader("Technical Analysis")
    ta_options = ["Moving Averages", "RSI", "MACD", "Bollinger Bands"]
    selected_ta = st.multiselect("Select Indicators", ta_options)
    
    # Market Sentiment
    st.subheader("Market Sentiment")
    sentiment_data = {
        'Source': ['News', 'Social Media', 'Analyst Ratings', 'Options Flow'],
        'Bullish': [65, 55, 70, 60],
        'Bearish': [35, 45, 30, 40]
    }
    fig = px.bar(pd.DataFrame(sentiment_data), x='Source', y=['Bullish', 'Bearish'],
                 title="Market Sentiment Analysis")
    st.plotly_chart(fig, use_container_width=True)
    
    # Economic Calendar
    st.subheader("Economic Calendar")
    calendar_data = {
        'Date': pd.date_range(start='2024-01-01', periods=5),
        'Event': ['Fed Rate Decision', 'GDP Release', 'CPI Data', 'Jobs Report', 'Retail Sales'],
        'Impact': ['High', 'High', 'Medium', 'High', 'Medium'],
        'Forecast': ['5.25%', '2.1%', '3.1%', '200K', '0.5%']
    }
    st.dataframe(pd.DataFrame(calendar_data))

elif page == "Optimization":
    st.title("Strategy Optimization")
    
    # Optimization Parameters
    st.subheader("Optimization Parameters")
    col1, col2 = st.columns(2)
    with col1:
        st.selectbox("Optimization Method", ["Genetic Algorithm", "Bayesian", "Grid Search"])
        st.number_input("Population Size", 10, 1000, 100)
    with col2:
        st.selectbox("Objective", ["Sharpe Ratio", "Returns", "Risk-Adjusted Returns"])
        st.number_input("Generations", 10, 100, 50)
    
    # Run Optimization
    if st.button("Run Optimization"):
        st.subheader("Optimization Results")
        
        # Parameter Space
        params_data = {
            'Parameter': ['Lookback Period', 'Entry Threshold', 'Exit Threshold', 'Position Size'],
            'Original': [20, 0.02, 0.01, 0.1],
            'Optimized': [25, 0.015, 0.008, 0.12]
        }
        st.dataframe(pd.DataFrame(params_data))
        
        # Performance Comparison
        dates = pd.date_range(start='2023-01-01', end='2024-01-01', freq='M')
        original_values = pd.Series(range(len(dates))) * 100 + 10000
        optimized_values = pd.Series(range(len(dates))) * 120 + 10000
        df_opt = pd.DataFrame({
            'Date': dates,
            'Original': original_values,
            'Optimized': optimized_values
        })
        fig = px.line(df_opt, x='Date', y=['Original', 'Optimized'],
                      title='Optimization Results')
        st.plotly_chart(fig, use_container_width=True)

elif page == "ML Models":
    st.title("Machine Learning Models")
    
    # Model Selection
    st.subheader("Model Management")
    model_options = ["LSTM", "Random Forest", "XGBoost", "Transformer"]
    selected_model = st.selectbox("Select Model", model_options)
    
    # Model Performance
    st.subheader("Model Performance")
    performance_data = {
        'Metric': ['Accuracy', 'Precision', 'Recall', 'F1-Score'],
        'Value': [0.85, 0.82, 0.88, 0.85]
    }
    st.dataframe(pd.DataFrame(performance_data))
    
    # Feature Importance
    st.subheader("Feature Importance")
    features = ['Price', 'Volume', 'RSI', 'MACD', 'Sentiment']
    importance = [0.3, 0.2, 0.15, 0.2, 0.15]
    
    fig = px.bar(x=features, y=importance, title="Feature Importance")
    st.plotly_chart(fig, use_container_width=True)
    
    # Model Training
    st.subheader("Model Training")
    if st.button("Train Model"):
        st.progress(100)
        st.success("Model training completed!")

elif page == "Settings":
    st.title("Settings")
    
    # API Configuration
    st.header("API Configuration")
    use_api = st.toggle("Enable API Features", value=st.session_state.use_api)
    if use_api:
        api_key = st.text_input("OpenAI API Key", value=st.session_state.api_key, type="password")
        if api_key != st.session_state.api_key:
            st.session_state.api_key = api_key
            st.session_state.llm = LLMInterface(api_key=api_key)
    st.session_state.use_api = use_api
    
    # Display current LLM status
    st.header("LLM Status")
    metrics = st.session_state.llm.get_llm_metrics()
    st.json(metrics)

    # Export/Import Configuration
    st.subheader("Export/Import Configuration")
    col1, col2 = st.columns(2)
    
    with col1:
        if st.button("Export Configuration"):
            config = {
                "use_api": st.session_state.use_api,
                "api_key": st.session_state.api_key
            }
            st.download_button(
                "Download Configuration",
                data=json.dumps(config, indent=2),
                file_name="trading_config.json",
                mime="application/json"
            )
    
    with col2:
        uploaded_file = st.file_uploader("Import Configuration", type=["json"])
        if uploaded_file is not None:
            try:
                config = json.load(uploaded_file)
                st.session_state.use_api = config.get("use_api", False)
                st.session_state.api_key = config.get("api_key", "")
                st.session_state.llm = LLMInterface(api_key=st.session_state.api_key)
                st.success("Configuration imported successfully!")
            except Exception as e:
                st.error(f"Error importing configuration: {str(e)}")

    # General Settings
    st.subheader("General Settings")
    st.selectbox("Theme", ["Light", "Dark", "System"])
    st.number_input("Refresh Rate (seconds)", 1, 60, 5)
    
    # Notification Settings
    st.subheader("Notifications")
    st.checkbox("Enable Email Notifications")
    st.checkbox("Enable SMS Notifications")
    st.checkbox("Enable Desktop Notifications")
    
    # Risk Settings
    st.subheader("Risk Parameters")
    st.number_input("Max Position Size (%)", 1, 100, 10)
    st.number_input("Max Drawdown (%)", 1, 50, 20)
    st.number_input("Daily Loss Limit (%)", 1, 20, 5) 

def create_performance_dashboard(symbol, data, market_data_df, reporting):
    """Create an interactive performance dashboard."""
    
    # Add export options
    st.sidebar.write("\nExport Options:")
    export_format = st.sidebar.selectbox(
        "Select Export Format",
        ["CSV", "Excel", "JSON", "Word", "PDF", "HTML"]
    )
    
    if st.sidebar.button("Export Data"):
        if export_format in ["CSV", "Excel", "JSON"]:
            export_data(data, format=export_format.lower(), filename=f"{symbol}_data")
        else:
            export_report(symbol, data, market_data_df, reporting, 
                         format=export_format.lower(), 
                         filename=f"{symbol}_report")
    
    # Create tabs for different sections
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "Performance Overview", 
        "Risk Analysis", 
        "Execution Metrics", 
        "Signal Analysis",
        "Regime Analysis"
    ])
    
    with tab1:
        # Performance Overview
        col1, col2 = st.columns(2)
        
        with col1:
            # Returns Distribution
            returns = data['Close'].pct_change().dropna()
            fig = go.Figure()
            fig.add_trace(go.Histogram(x=returns, nbinsx=50, name='Returns'))
            fig.add_trace(go.Scatter(x=[returns.mean()]*2, y=[0, 100], 
                                   mode='lines', name='Mean',
                                   line=dict(color='red', dash='dash')))
            fig.update_layout(title='Returns Distribution',
                            xaxis_title='Return',
                            yaxis_title='Frequency',
                            showlegend=True)
            st.plotly_chart(fig)
            
            # Cumulative Returns
            cum_returns = (1 + returns).cumprod()
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=cum_returns.index, y=cum_returns,
                                   mode='lines', name='Cumulative Returns'))
            fig.update_layout(title='Cumulative Returns',
                            xaxis_title='Date',
                            yaxis_title='Cumulative Return',
                            showlegend=True)
            st.plotly_chart(fig)
        
        with col2:
            # Performance Metrics
            st.metric("Total Return", f"{reporting['performance_summary']['total_return']:.2%}")
            st.metric("Annualized Return", f"{reporting['performance_summary']['annualized_return']:.2%}")
            st.metric("Sharpe Ratio", f"{reporting['performance_summary']['sharpe_ratio']:.2f}")
            st.metric("Sortino Ratio", f"{reporting['performance_summary']['sortino_ratio']:.2f}")
            st.metric("Max Drawdown", f"{reporting['performance_summary']['max_drawdown']:.2%}")
            st.metric("Win Rate", f"{reporting['performance_summary']['win_rate']:.2%}")
            st.metric("Profit Factor", f"{reporting['performance_summary']['profit_factor']:.2f}")
    
    with tab2:
        # Risk Analysis
        col1, col2 = st.columns(2)
        
        with col1:
            # Risk Metrics
            st.metric("VaR (95%)", f"{reporting['risk_metrics']['var_95']:.2%}")
            st.metric("Expected Shortfall (95%)", f"{reporting['risk_metrics']['expected_shortfall_95']:.2%}")
            st.metric("Beta", f"{reporting['risk_metrics']['beta']:.2f}")
            st.metric("Correlation", f"{reporting['risk_metrics']['correlation']:.2f}")
            st.metric("Tracking Error", f"{reporting['risk_metrics']['tracking_error']:.2%}")
            st.metric("Information Ratio", f"{reporting['risk_metrics']['information_ratio']:.2f}")
        
        with col2:
            # Risk Visualization
            # Rolling Volatility
            rolling_vol = returns.rolling(window=20).std() * np.sqrt(252)
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=rolling_vol.index, y=rolling_vol,
                                   mode='lines', name='Rolling Volatility'))
            fig.update_layout(title='Rolling Volatility (20-day)',
                            xaxis_title='Date',
                            yaxis_title='Volatility',
                            showlegend=True)
            st.plotly_chart(fig)
            
            # Drawdown Analysis
            cummax = cum_returns.cummax()
            drawdown = (cum_returns - cummax) / cummax
            fig = go.Figure()
            fig.add_trace(go.Scatter(x=drawdown.index, y=drawdown,
                                   mode='lines', name='Drawdown',
                                   fill='tozeroy'))
            fig.update_layout(title='Drawdown Analysis',
                            xaxis_title='Date',
                            yaxis_title='Drawdown',
                            showlegend=True)
            st.plotly_chart(fig)
    
    with tab3:
        # Execution Metrics
        col1, col2 = st.columns(2)
        
        with col1:
            # Execution Metrics
            st.metric("Total Orders", reporting['execution_metrics']['total_orders'])
            st.metric("Filled Orders", reporting['execution_metrics']['filled_orders'])
            st.metric("Cancelled Orders", reporting['execution_metrics']['cancelled_orders'])
            st.metric("Average Slippage", f"{reporting['execution_metrics']['average_slippage']:.4%}")
            st.metric("Average Execution Time", f"{reporting['execution_metrics']['average_execution_time']:.1f}s")
            st.metric("Average Price Improvement", f"{reporting['execution_metrics']['average_price_improvement']:.4%}")
            st.metric("Average Market Impact", f"{reporting['execution_metrics']['average_market_impact']:.4%}")
            st.metric("Execution Cost", f"{reporting['execution_metrics']['execution_cost']:.4%}")
        
        with col2:
            # Execution Visualization
            # Slippage Distribution
            if 'slippage' in reporting['execution_metrics']:
                slippage = reporting['execution_metrics']['slippage']
                fig = go.Figure()
                fig.add_trace(go.Histogram(x=slippage, nbinsx=50, name='Slippage'))
                fig.update_layout(title='Slippage Distribution',
                                xaxis_title='Slippage',
                                yaxis_title='Frequency',
                                showlegend=True)
                st.plotly_chart(fig)
            
            # Execution Time Distribution
            if 'execution_time' in reporting['execution_metrics']:
                exec_time = reporting['execution_metrics']['execution_time']
                fig = go.Figure()
                fig.add_trace(go.Histogram(x=exec_time, nbinsx=50, name='Execution Time'))
                fig.update_layout(title='Execution Time Distribution',
                                xaxis_title='Time (s)',
                                yaxis_title='Frequency',
                                showlegend=True)
                st.plotly_chart(fig)
    
    with tab4:
        # Signal Analysis
        col1, col2 = st.columns(2)
        
        with col1:
            # Signal Metrics
            st.metric("Total Signals", reporting['signal_metrics']['total_signals'])
            st.metric("Accurate Signals", reporting['signal_metrics']['accurate_signals'])
            st.metric("Win Rate", f"{reporting['signal_metrics']['win_rate']:.2%}")
            st.metric("Profit Factor", f"{reporting['signal_metrics']['profit_factor']:.2f}")
            st.metric("Average Return", f"{reporting['signal_metrics']['average_return']:.2%}")
            st.metric("Sharpe Ratio", f"{reporting['signal_metrics']['sharpe_ratio']:.2f}")
        
        with col2:
            # Signal Visualization
            # Signal Performance Over Time
            if 'signal_performance' in reporting:
                signal_perf = reporting['signal_performance']
                fig = go.Figure()
                fig.add_trace(go.Scatter(x=signal_perf['dates'], 
                                       y=signal_perf['returns'],
                                       mode='lines', name='Signal Returns'))
                fig.update_layout(title='Signal Performance Over Time',
                                xaxis_title='Date',
                                yaxis_title='Return',
                                showlegend=True)
                st.plotly_chart(fig)
            
            # Signal Accuracy Distribution
            if 'signal_accuracy' in reporting:
                accuracy = reporting['signal_accuracy']
                fig = go.Figure()
                fig.add_trace(go.Histogram(x=accuracy, nbinsx=50, name='Signal Accuracy'))
                fig.update_layout(title='Signal Accuracy Distribution',
                                xaxis_title='Accuracy',
                                yaxis_title='Frequency',
                                showlegend=True)
                st.plotly_chart(fig)
    
    with tab5:
        # Regime Analysis
        col1, col2 = st.columns(2)
        
        with col1:
            # Regime Metrics
            st.metric("Current Regime", reporting['regime_metrics']['current_regime'])
            st.metric("Regime Duration", f"{reporting['regime_metrics']['regime_duration']:.1f}h")
            st.metric("Regime Return", f"{reporting['regime_metrics']['regime_return']:.2%}")
            st.metric("Regime Volatility", f"{reporting['regime_metrics']['regime_volatility']:.2%}")
            st.metric("Regime Sharpe", f"{reporting['regime_metrics']['regime_sharpe']:.2f}")
        
        with col2:
            # Regime Visualization
            # Regime Transitions
            if 'regime_transitions' in reporting:
                transitions = reporting['regime_transitions']
                fig = go.Figure()
                fig.add_trace(go.Scatter(x=transitions['dates'],
                                       y=transitions['regimes'],
                                       mode='lines+markers',
                                       name='Regime Transitions'))
                fig.update_layout(title='Regime Transitions',
                                xaxis_title='Date',
                                yaxis_title='Regime',
                                showlegend=True)
                st.plotly_chart(fig)
            
            # Regime Performance
            if 'regime_performance' in reporting:
                regime_perf = reporting['regime_performance']
                fig = go.Figure()
                for regime, returns in regime_perf.items():
                    fig.add_trace(go.Box(y=returns,
                                       name=regime,
                                       boxpoints='all'))
                fig.update_layout(title='Regime Performance Distribution',
                                xaxis_title='Regime',
                                yaxis_title='Returns',
                                showlegend=True)
                st.plotly_chart(fig)

# ... existing code ...

            elif action['type'] == 'analyze_market':
                symbol = action.get('symbol', 'AAPL')
                market_symbol = action.get('market_symbol', '^GSPC')  # S&P 500 as default market
                
                # Get historical data
                data = st.session_state.market_data.get_historical_data(symbol, window=100)
                market_data_df = st.session_state.market_data.get_historical_data(market_symbol, window=100)
                
                if data is not None and not data.empty and market_data_df is not None and not market_data_df.empty:
                    # ... existing market analysis code ...
                    
                    # Create interactive performance dashboard
                    create_performance_dashboard(symbol, data, market_data_df, st.session_state.reporting)
                    
                    # Create enhanced visualization
                    fig = make_subplots(rows=23, cols=1,
                                      shared_xaxes=True,
                                      vertical_spacing=0.05,
                                      subplot_titles=('Price and Trend', 'Volatility', 'Correlation', 
                                                    'News Sentiment', 'Entity Mentions', 'Trading Signals',
                                                    'Signal Performance', 'Market Filters', 'Market Regime',
                                                    'Regime Transitions', 'Signal Patterns', 'Risk Metrics',
                                                    'Execution Metrics', 'Position Management', 'Portfolio Risk',
                                                    'Portfolio Optimization', 'Execution Management',
                                                    'Market Impact', 'Adaptive Execution', 'Performance Analytics',
                                                    'Portfolio Performance', 'Performance Summary',
                                                    'Risk Summary'),
                                      row_heights=[0.2, 0.15, 0.15, 0.15, 0.15, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1, 0.1])
                    
                    # ... existing plots ...
                    
                    # Performance summary plot
                    performance_summary = reporting['performance_summary']
                    metrics = ['Total Return', 'Largest Win', 'Largest Loss']
                    values = [performance_summary['total_return'], 
                            performance_summary['largest_win'],
                            performance_summary['largest_loss']]
                    
                    fig.add_trace(go.Bar(x=metrics, y=values,
                                       name='Performance Summary',
                                       marker_color=['blue', 'green', 'red']),
                                row=22, col=1)
                    
                    # Risk summary plot
                    risk_metrics = reporting['risk_metrics']
                    metrics = ['VaR (95%)', 'Expected Shortfall', 'Beta', 'Correlation']
                    values = [risk_metrics['var_95'],
                            risk_metrics['expected_shortfall_95'],
                            risk_metrics['beta'],
                            risk_metrics['correlation']]
                    
                    fig.add_trace(go.Bar(x=metrics, y=values,
                                       name='Risk Summary',
                                       marker_color=['purple', 'orange', 'cyan', 'magenta']),
                                row=23, col=1)
                    
                    fig.update_layout(height=4200, title_text=f'Enhanced Market Analysis for {symbol}')
                    st.plotly_chart(fig)

# ... rest of the code ...

def generate_report(symbol, data, market_data_df, reporting):
    """Generate a comprehensive report with visualizations."""
    
    # Create report template
    template = """
    <!DOCTYPE html>
    <html>
    <head>
        <title>Trading System Report - {{ symbol }}</title>
        <style>
            body { font-family: Arial, sans-serif; margin: 40px; }
            h1 { color: #2c3e50; }
            h2 { color: #34495e; margin-top: 30px; }
            .metric { margin: 10px 0; }
            .metric-label { font-weight: bold; }
            .metric-value { color: #2980b9; }
            .chart { margin: 20px 0; }
            table { border-collapse: collapse; width: 100%; margin: 20px 0; }
            th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }
            th { background-color: #f5f5f5; }
            .positive { color: #27ae60; }
            .negative { color: #c0392b; }
        </style>
    </head>
    <body>
        <h1>Trading System Report - {{ symbol }}</h1>
        <p>Generated on: {{ timestamp }}</p>
        
        <h2>Performance Summary</h2>
        <div class="metrics">
            {% for metric, value in performance.items() %}
            <div class="metric">
                <span class="metric-label">{{ metric|replace('_', ' ')|title }}:</span>
                <span class="metric-value">{{ value|format_value }}</span>
            </div>
            {% endfor %}
        </div>
        
        <h2>Risk Metrics</h2>
        <div class="metrics">
            {% for metric, value in risk.items() %}
            <div class="metric">
                <span class="metric-label">{{ metric|replace('_', ' ')|title }}:</span>
                <span class="metric-value">{{ value|format_value }}</span>
            </div>
            {% endfor %}
        </div>
        
        <h2>Execution Metrics</h2>
        <div class="metrics">
            {% for metric, value in execution.items() %}
            <div class="metric">
                <span class="metric-label">{{ metric|replace('_', ' ')|title }}:</span>
                <span class="metric-value">{{ value|format_value }}</span>
            </div>
            {% endfor %}
        </div>
        
        <h2>Signal Metrics</h2>
        <div class="metrics">
            {% for metric, value in signals.items() %}
            <div class="metric">
                <span class="metric-label">{{ metric|replace('_', ' ')|title }}:</span>
                <span class="metric-value">{{ value|format_value }}</span>
            </div>
            {% endfor %}
        </div>
        
        <h2>Regime Metrics</h2>
        <div class="metrics">
            {% for metric, value in regime.items() %}
            <div class="metric">
                <span class="metric-label">{{ metric|replace('_', ' ')|title }}:</span>
                <span class="metric-value">{{ value|format_value }}</span>
            </div>
            {% endfor %}
        </div>
        
        <h2>Charts</h2>
        <div class="charts">
            {{ charts|safe }}
        </div>
    </body>
    </html>
    """
    
    # Create Jinja2 environment
    env = jinja2.Environment()
    env.filters['format_value'] = lambda x: f"{x:.2%}" if isinstance(x, float) else str(x)
    
    # Generate charts
    charts = []
    
    # Returns Distribution
    returns = data['Close'].pct_change().dropna()
    fig = go.Figure()
    fig.add_trace(go.Histogram(x=returns, nbinsx=50, name='Returns'))
    fig.add_trace(go.Scatter(x=[returns.mean()]*2, y=[0, 100], 
                           mode='lines', name='Mean',
                           line=dict(color='red', dash='dash')))
    fig.update_layout(title='Returns Distribution',
                    xaxis_title='Return',
                    yaxis_title='Frequency',
                    showlegend=True)
    charts.append(fig.to_html(full_html=False))
    
    # Cumulative Returns
    cum_returns = (1 + returns).cumprod()
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=cum_returns.index, y=cum_returns,
                           mode='lines', name='Cumulative Returns'))
    fig.update_layout(title='Cumulative Returns',
                    xaxis_title='Date',
                    yaxis_title='Cumulative Return',
                    showlegend=True)
    charts.append(fig.to_html(full_html=False))
    
    # Rolling Volatility
    rolling_vol = returns.rolling(window=20).std() * np.sqrt(252)
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=rolling_vol.index, y=rolling_vol,
                           mode='lines', name='Rolling Volatility'))
    fig.update_layout(title='Rolling Volatility (20-day)',
                    xaxis_title='Date',
                    yaxis_title='Volatility',
                    showlegend=True)
    charts.append(fig.to_html(full_html=False))
    
    # Drawdown Analysis
    cummax = cum_returns.cummax()
    drawdown = (cum_returns - cummax) / cummax
    fig = go.Figure()
    fig.add_trace(go.Scatter(x=drawdown.index, y=drawdown,
                           mode='lines', name='Drawdown',
                           fill='tozeroy'))
    fig.update_layout(title='Drawdown Analysis',
                    xaxis_title='Date',
                    yaxis_title='Drawdown',
                    showlegend=True)
    charts.append(fig.to_html(full_html=False))
    
    # Render template
    template = env.from_string(template)
    html = template.render(
        symbol=symbol,
        timestamp=datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        performance=reporting['performance_summary'],
        risk=reporting['risk_metrics'],
        execution=reporting['execution_metrics'],
        signals=reporting['signal_metrics'],
        regime=reporting['regime_metrics'],
        charts='\n'.join(charts)
    )
    
    return html

# ... rest of the code ...

def export_report(symbol, data, market_data_df, reporting, format='pdf', filename=None):
    """Export comprehensive report to various formats."""
    if filename is None:
        filename = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    
    # Generate HTML report
    html = generate_report(symbol, data, market_data_df, reporting)
    
    # Export based on format
    if format == 'html':
        st.download_button(
            label="Download HTML Report",
            data=html,
            file_name=f"{filename}.html",
            mime="text/html"
        )
    elif format == 'pdf':
        pdf = pdfkit.from_string(html, False)
        st.download_button(
            label="Download PDF Report",
            data=pdf,
            file_name=f"{filename}.pdf",
            mime="application/pdf"
        )
    elif format == 'docx':
        doc = Document()
        doc.add_heading('Trading System Report', 0)
        
        # Add timestamp
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Add sections
        sections = [
            ('Performance Summary', reporting['performance_summary']),
            ('Risk Metrics', reporting['risk_metrics']),
            ('Execution Metrics', reporting['execution_metrics']),
            ('Signal Metrics', reporting['signal_metrics']),
            ('Regime Metrics', reporting['regime_metrics'])
        ]
        
        for title, metrics in sections:
            doc.add_heading(title, level=1)
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = 'Metric'
            table.rows[0].cells[1].text = 'Value'
            
            for metric, value in metrics.items():
                if metric != 'last_update':
                    row = table.add_row()
                    row.cells[0].text = metric.replace('_', ' ').title()
                    row.cells[1].text = f"{value:.2%}" if isinstance(value, float) else str(value)
        
        # Save and provide download
        doc.save(f"{filename}.docx")
        with open(f"{filename}.docx", "rb") as file:
            st.download_button(
                label="Download Word Document",
                data=file,
                file_name=f"{filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

# ... rest of the code ...